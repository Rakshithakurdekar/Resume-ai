import os
import re
import io
import json
import time
import base64
import PyPDF2
import nltk
import pandas as pd
import streamlit as st
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import plotly.graph_objects as go
from PIL import Image

# Optional Library Checks
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except Exception:
    FPDF_AVAILABLE = False
    
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# -------------------------
# NLTK setup (remains the same)
# -------------------------
def setup_nltk():
    try:
        sw = set(stopwords.words('english'))
    except LookupError:
        nltk.download('stopwords', quiet=True)
        sw = set(stopwords.words('english'))

    try:
        nltk.data.find("corpora/wordnet")
    except LookupError:
        nltk.download("wordnet", quiet=True)

    try:
        nltk.data.find("tokenizers/punkt")
    except LookupError:
        nltk.download("punkt", quiet=True)

    return frozenset(sw), WordNetLemmatizer()


STOPWORDS, lemmatizer = setup_nltk()


# -------------------------
# Skill lists (remains the same)
# -------------------------
PREMIER_SKILLS = {
    "Data Science": ["python", "r", "sql", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch",
                     "machine learning", "deep learning", "nlp", "statistics"],

    "Engineering": ["java", "c++", "javascript", "react", "node.js", "docker", "aws", "kubernetes", "git"],

    "Business": ["excel", "tableau", "power bi", "communication", "leadership"],
}
ALL_SKILLS = [s for skills in PREMIER_SKILLS.values() for s in skills]


# -------------------------
# Text utilities (remains the same)
# -------------------------
def extract_text_from_pdf(file):
    text = ""
    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            p = page.extract_text()
            if p:
                text += p + "\n"
    except Exception:
        return ""
    return text


def preprocess(text: str) -> str:
    if not text:
        return ""

    text = text.lower()
    text = re.sub(r"[^a-z0-9\s]", " ", text)

    tokens = []
    for w in text.split():
        if w not in STOPWORDS:
            tokens.append(lemmatizer.lemmatize(w))
    return " ".join(tokens)


def advanced_tfidf_similarity(resume: str, job: str):
    resume = resume or ""
    job = job or ""

    try:
        vectorizer = TfidfVectorizer(ngram_range=(1,3), max_features=7000)
        tfidf = vectorizer.fit_transform([resume, job])
        score = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
        return round(score * 100, 2), vectorizer, tfidf
    except:
        return 0.0, None, None


def extract_skills(text: str):
    text = (text or "").lower()
    return [s for s in ALL_SKILLS if s in text]


def compare_skills(resume_text: str, job_text: str):
    job = set(extract_skills(job_text))
    res = set(extract_skills(resume_text))
    return sorted(list(job & res)), sorted(list(job - res))


def generate_ai_insights(overall, tfidf, missing_skills, missing_kw):
    insights = []

    if overall >= 85:
        insights.append("⭐ Excellent match! Very strong fit.")
    elif overall >= 70:
        insights.append("👍 Good match. Small improvements can increase your score.")
    else:
        insights.append("⚠ Resume needs improvements to match JD better.")

    if missing_skills:
        insights += [f"💡 Add or highlight experience in: {s}" for s in missing_skills[:5]]

    if missing_kw:
        insights.append("📝 Add important keywords: " + ", ".join(missing_kw[:10]))

    return insights


def create_report(name, tfidf, skill, overall, matched_kw, missing_kw, matched_sk, missing_sk, insights, filename):
    return {
        "candidate": name,
        "file": filename,
        "tfidf_score": tfidf,
        "skill_match_percent": skill,
        "overall_score": overall,
        "matched_keywords": matched_kw,
        "missing_keywords": missing_kw,
        "matched_skills": matched_sk,
        "missing_skills": missing_sk,
        "ai_insights": insights,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
    }


# -------------------------
# Visualization helpers (remains the same)
# -------------------------
def plot_skill_radar(matched_skills, missing_skills):
    categories = list(set(matched_skills + missing_skills))
    if not categories:
        return None

    # Using 1 for matched, 0.2 for missing/not present
    matched_vals = [1 if s in matched_skills else 0.2 for s in categories]

    categories += [categories[0]] # Close the loop
    matched_vals += [matched_vals[0]]

    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=matched_vals, theta=categories, fill='toself', name='Skill Match'))
    fig.update_layout(
        template="plotly_dark",
        showlegend=False,
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 1])
        )
    )
    return fig


def keyword_heatmap(jd_tokens, resume_tokens):
    jd_freq = {}
    for t in jd_tokens:
        jd_freq[t] = jd_freq.get(t, 0) + 1

    top = sorted(jd_freq.items(), key=lambda x: -x[1])[:20]
    words = [w for w, _ in top]
    presence = [1 if w in resume_tokens else 0 for w in words]

    # Create a simple bar chart to show keyword presence
    fig = go.Figure(go.Bar(
        x=words,
        y=presence,
        marker_color=['#1f77b4' if p == 1 else '#d62728' for p in presence] # Blue for present, Red for missing
    ))
    fig.update_layout(
        title="Top 20 JD Keywords Match",
        yaxis_title="Present in Resume (1=Yes, 0=No)",
        template="plotly_dark"
    )
    return fig

# ---------------------------------
# NEW: PDF Generation Function
# ---------------------------------
def generate_pdf_report(report_data):
    if not FPDF_AVAILABLE:
        st.error("FPDF library not installed. Cannot generate PDF.")
        return None

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.set_text_color(0, 50, 100)
    pdf.cell(0, 10, "AI Resume Matcher Report", 0, 1, "C")
    pdf.set_font("Arial", "", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 5, f"Candidate File: {report_data['file']}", 0, 1)
    pdf.cell(0, 5, f"Analysis Date: {report_data['timestamp']}", 0, 1)
    pdf.ln(5)

    # Scores
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 7, "Match Scores", 0, 1, "L")
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 6, f"Overall Score: {report_data['overall_score']}%", 0, 1)
    pdf.cell(0, 6, f"TF-IDF Semantic Match: {report_data['tfidf_score']}%", 0, 1)
    pdf.cell(0, 6, f"Extracted Skill Match: {report_data['skill_match_percent']}%", 0, 1)
    pdf.ln(5)

    # Insights
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 7, "AI Insights & Recommendations", 0, 1, "L")
    pdf.set_font("Arial", "", 12)
    for insight in report_data['ai_insights']:
        pdf.multi_cell(0, 5, f"- {insight}")
    pdf.ln(5)

    # Keywords/Skills (simplified for PDF)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 7, "Keywords and Skills Breakdown", 0, 1, "L")
    pdf.set_font("Arial", "", 12)

    pdf.cell(0, 6, "Matched Keywords:", 0, 1)
    pdf.multi_cell(0, 5, ", ".join(report_data['matched_keywords'][:20]) or "None")
    
    pdf.cell(0, 6, "Missing Keywords (Top 20):", 0, 1)
    pdf.multi_cell(0, 5, ", ".join(report_data['missing_keywords'][:20]) or "None")

    pdf.cell(0, 6, "Missing Skills:", 0, 1)
    pdf.multi_cell(0, 5, ", ".join(report_data['missing_skills']) or "None")

    # Output buffer
    pdf_output = pdf.output(dest='S').encode('latin1')
    return pdf_output

# ---------------------------------
# NEW: DOCX Generation Function
# ---------------------------------
def generate_docx_report(report_data):
    if not DOCX_AVAILABLE:
        st.error("python-docx library not installed. Cannot generate Word document.")
        return None
        
    document = Document()
    document.add_heading('AI Resume Matcher Report', 0)
    document.add_paragraph(f"Candidate File: {report_data['file']}")
    document.add_paragraph(f"Analysis Date: {report_data['timestamp']}")
    
    document.add_heading('Match Scores', 1)
    document.add_paragraph(f"Overall Score: {report_data['overall_score']}%")
    document.add_paragraph(f"TF-IDF Semantic Match: {report_data['tfidf_score']}%")
    document.add_paragraph(f"Extracted Skill Match: {report_data['skill_match_percent']}%")

    document.add_heading('AI Insights & Recommendations', 1)
    for insight in report_data['ai_insights']:
        document.add_paragraph(insight, style='List Bullet')

    document.add_heading('Keywords and Skills Breakdown', 1)

    document.add_heading('Matched Keywords', 2)
    document.add_paragraph(", ".join(report_data['matched_keywords']) or "N/A")

    document.add_heading('Missing Keywords', 2)
    document.add_paragraph(", ".join(report_data['missing_keywords']) or "N/A")
    
    document.add_heading('Missing Skills', 2)
    document.add_paragraph(", ".join(report_data['missing_skills']) or "N/A")

    # Save to a temporary buffer
    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.read()

# ======================================================
#                LANDING PAGE + ABOUT (remains the same)
# ======================================================

# Placeholder image creation 
try:
    HERO_IMAGE = Image.open("image .png")
    RESUME_ICON = Image.open("image copy 2.png")
except FileNotFoundError:
    HERO_IMAGE = Image.new('RGB', (50, 50), color = '#3498db') 
    RESUME_ICON = Image.new('RGB', (100, 100), color = '#2ecc71')
    st.warning("⚠️ Using placeholder images. Please ensure 'image copy.png' and 'image copy 2.png' are in the directory.")


def landing_page():
    # --- Hero Section ---
    st.image(HERO_IMAGE, use_container_width=True)
    
    st.markdown("""
        <h1 style='text-align:center;'>Empower Your Career with AI Intelligence</h1>
        <h3 style='text-align:center;'>Instant resume analysis, skill gap detection, and AI recommendations</h3>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🚀 START AI RESUME ANALYSIS 🚀", key="start_hero", use_container_width=True):
            st.session_state["page"] = "matcher"
    
    with col3:
        st.button("Learn More 💡", key="learn_hero", use_container_width=False)

    st.markdown("---")

    # --- Feature Boxes ---
    st.markdown("## 🌟 Key Features")
    cols = st.columns(4)
    features = [
        ("Resume Improvement", "Helps job seekers improve resumes with AI suggestions.", RESUME_ICON),
        ("Fast Resume Filtering", "Recruiters can filter and analyse resumes faster.", RESUME_ICON),
        ("Skill Gap Detection", "Students and candidates understand which skills they lack.", RESUME_ICON),
        ("Job Matching", "AI helps companies match the right candidates efficiently.", RESUME_ICON),
    ]
    
    for col, feature in zip(cols, features):
        with col:
            st.image(feature[2], width=120) 
            st.markdown(f"**{feature[0]}**")
            st.markdown(feature[1])

    st.markdown("---")

    # --- About/How it Works Section ---
    st.markdown("## ⭐ About Us & Core Workflow")
    st.markdown("""
    The **AI Resume Talent Scout** is an intelligent system created to help:
    - Job seekers improve their resumes  
    - Recruiters filter and analyse resumes faster  
    - Students understand skills they lack  
    - Companies match the right candidates  
    """)
    
    st.markdown("### 🔍 Core Workflow")
    st.markdown("""
    - Extracts text from resumes  
    - Compares with Job Description  
    - Uses AI (**TF-IDF** + NLP) to compute similarity  
    - Detects matched & missing skills  
    - Generates insights & recommendations  
    """)

    st.markdown("### 🛠 Problems Solved")
    st.markdown("""
    - Missing keywords in resumes
    - No measurable matching score
    - Lack of instant skill gap identification
    - Manual shortlisting
    """)

    st.markdown("### 🚀 Enhancements Included")
    st.markdown("""
    - TF-IDF semantic similarity 
    - Skill radar chart 
    - Keyword heatmap 
    - AI recommendations 
    - Professional PDF report
    - Clean UI with dark theme support 
    """)

    st.markdown("---")
    col_b1, col_b2, col_b3 = st.columns([1, 2, 1])
    with col_b2:
        if st.button("START RESUME MATCHING NOW! 🚀", key="start_bottom", use_container_width=True):
            st.session_state["page"] = "matcher"


# ======================================================
#                  RESUME MATCHER PAGE
# ======================================================

def matcher_page():
    st.title("Resume Matcher")
    
    if st.button("⬅️ Back to Home"):
        st.session_state["page"] = "landing"
        st.rerun()
        return

    with st.sidebar:
        st.header("Settings")
        candidate_prefix = st.text_input("Candidate name", "Candidate")
        allow_multiple = st.checkbox("Upload multiple resumes", value=True)

    resumes = st.file_uploader("Upload Resume PDF(s)", type=["pdf"], accept_multiple_files=allow_multiple)
    jd = st.text_area("Paste Job Description", height=220)

    run = st.button("Run Analysis")

    if run:
        if not resumes or not jd.strip():
            st.error("Upload resume and paste JD first!")
            return

        jd_proc = preprocess(jd)
        results = []
        
        # Display progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for idx, f in enumerate(resumes):
            status_text.text(f"Processing {f.name}...")
            
            raw_text = extract_text_from_pdf(f)
            proc_text = preprocess(raw_text)
            
            # TF-IDF calculation
            tfidf_score, _, _ = advanced_tfidf_similarity(proc_text, jd_proc)

            resume_tokens = set(proc_text.split())
            jd_tokens = set(jd_proc.split())

            matched_kw = sorted(list(resume_tokens & jd_tokens))
            missing_kw = sorted(list(jd_tokens - resume_tokens))

            # Skill comparison
            matched_sk, missing_sk = compare_skills(proc_text, jd_proc)

            skill_score = round(len(matched_sk) / max(1, len(ALL_SKILLS)) * 100, 2)
            
            overall = round(0.6 * tfidf_score + 0.4 * skill_score, 2)

            insights = generate_ai_insights(overall, tfidf_score, missing_sk, missing_kw)

            report = create_report(
                candidate_prefix, tfidf_score, skill_score, overall,
                matched_kw, missing_kw, matched_sk, missing_sk,
                insights, f.name
            )

            results.append({"file": f.name, "report": report, "raw_text": raw_text, "proc_text": proc_text})
            progress_bar.progress((idx + 1) / len(resumes))
        
        status_text.success("Analysis Completed!")
        progress_bar.empty()

        st.markdown("## 📊 Analysis Results")

        for item in results:
            rep = item["report"]
            st.markdown("---")
            st.header(f"Results for: {item['file']}")

            # Score Gauge
            st.subheader(f"Overall Match Score: {rep['overall_score']}%")
            
            col_sc1, col_sc2 = st.columns(2)
            col_sc1.metric("TF-IDF Semantic Match", f"{rep['tfidf_score']}%")
            col_sc2.metric("Extracted Skill Match", f"{rep['skill_match_percent']}%")

            # Insights
            st.subheader("💡 AI Insights & Recommendations")
            for i in rep["ai_insights"]:
                st.write(f"- {i}")
            
            # Detailed Breakdown
            with st.expander("Detailed Breakdown and Visualizations"):
                
                # Charts
                st.subheader("Keyword Heatmap (Presence in JD vs Resume)")
                st.plotly_chart(keyword_heatmap(jd_tokens, resume_tokens), use_container_width=True)
                
                st.subheader("Skill Match Radar")
                st.plotly_chart(plot_skill_radar(rep["matched_skills"], rep["missing_skills"]), use_container_width=True)

                st.subheader("Matched Keywords (from JD)")
                st.info(", ".join(rep["matched_keywords"]))

                st.subheader("Missing Keywords (from JD)")
                st.warning(", ".join(rep["missing_keywords"]))

                st.subheader("Matched Skills (from predefined list)")
                st.success(", ".join(rep["matched_skills"]))

                st.subheader("Missing Skills (from predefined list)")
                st.error(", ".join(rep["missing_skills"]))

            # --- Download Buttons ---
            st.markdown("### ⬇️ Download Report")
            col_down1, col_down2, col_down3 = st.columns(3)

            # Download JSON
            with col_down1:
                st.download_button(
                    "Download JSON",
                    json.dumps(rep, indent=2).encode(),
                    file_name=f"{rep['file']}_report.json",
                    use_container_width=True
                )

            # Download PDF
            if FPDF_AVAILABLE:
                with col_down2:
                    pdf_data = generate_pdf_report(rep)
                    st.download_button(
                        "Download PDF Report",
                        data=pdf_data,
                        file_name=f"{rep['file']}_report.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
            else:
                col_down2.warning("Install fpdf to enable PDF download.")

            # Download Word Document (DOCX)
            if DOCX_AVAILABLE:
                with col_down3:
                    docx_data = generate_docx_report(rep)
                    st.download_button(
                        "Download DOCX Report",
                        data=docx_data,
                        file_name=f"{rep['file']}_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                col_down3.warning("Install python-docx to enable DOCX download.")


# ======================================================
# MAIN APP ROUTER (remains the same)
# ======================================================

def main():
    st.set_page_config(page_title="AI Resume Matcher", layout="wide")

    if "page" not in st.session_state:
        st.session_state["page"] = "landing"

    if st.session_state["page"] == "landing":
        landing_page()
    else:
        matcher_page()


if __name__ == "__main__":
    main()